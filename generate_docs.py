import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure directories exist
os.makedirs('public/docs', exist_ok=True)

# --- PDF Presentation Generation ---
def create_presentation():
    doc = SimpleDocTemplate("public/demo-presentation.pdf", pagesize=landscape(A4),
                            rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=36,
        textColor=colors.HexColor("#D4A853"), # Accent color
        alignment=TA_CENTER,
        spaceAfter=30
    )
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Heading2'],
        fontSize=24,
        textColor=colors.white,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        fontSize=16,
        textColor=colors.white,
        leading=24,
        alignment=TA_LEFT
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['BodyText'],
        fontSize=16,
        textColor=colors.white,
        leading=24,
        leftIndent=20,
        bulletIndent=10
    )

    # Background color function
    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#0A0A0A")) # Dark background
        canvas.rect(0, 0, landscape(A4)[0], landscape(A4)[1], fill=1)
        # Add footer
        canvas.setFillColor(colors.HexColor("#D4A853"))
        canvas.setFont("Helvetica", 10)
        canvas.drawString(30, 20, "Digni Digital - Future-Ready Graduate Program")
        canvas.drawString(landscape(A4)[0]-50, 20, str(doc.page))
        canvas.restoreState()

    story = []

    # Slide 1: Title
    story.append(Spacer(1, 100))
    story.append(Paragraph("Future-Ready Graduate Program", title_style))
    story.append(Paragraph("Transforming Education for the Digital Age", subtitle_style))
    story.append(Spacer(1, 50))
    story.append(Paragraph("Prepared for: Forward-Thinking Educational Institutions", 
                           ParagraphStyle('CenterBody', parent=body_style, alignment=TA_CENTER)))
    story.append(PageBreak())

    # Slide 2: The Challenge
    story.append(Paragraph("The Challenge", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Traditional education is facing a crisis of relevance:", body_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("• 40% of graduates remain unemployed after 6 months", bullet_style))
    story.append(Paragraph("• 75% of employers say graduates lack job-ready skills", bullet_style))
    story.append(Paragraph("• Students graduate with theory but no income-generating skills", bullet_style))
    story.append(Paragraph("• The 'Digital Divide' is leaving students behind", bullet_style))
    story.append(PageBreak())

    # Slide 3: Our Solution
    story.append(Paragraph("The Solution: Future-Ready Graduate™", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("We transform your final-year students into skilled professionals.", body_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("An all-in-one partnership that brings:", body_style))
    story.append(Paragraph("• Practical Digital Skills Training (Web Dev, Marketing, AI)", bullet_style))
    story.append(Paragraph("• Real-world Portfolio Development", bullet_style))
    story.append(Paragraph("• Direct Employer Connections", bullet_style))
    story.append(Paragraph("• Seamless Integration with Academic Calendar", bullet_style))
    story.append(PageBreak())

    # Slide 4: Partnership Model
    story.append(Paragraph("How It Works", title_style))
    story.append(Spacer(1, 20))
    
    data = [
        [Paragraph("<b>What We Provide</b>", body_style), Paragraph("<b>What You Provide</b>", body_style)],
        [Paragraph("• High-Speed Internet Connectivity", bullet_style), Paragraph("• Computer Lab Facility", bullet_style)],
        [Paragraph("• Expert On-site Facilitators", bullet_style), Paragraph("• Reliable Electricity", bullet_style)],
        [Paragraph("• Premium AI Tools (ChatGPT, etc.)", bullet_style), Paragraph("• Final Year Students", bullet_style)],
        [Paragraph("• Complete Curriculum", bullet_style), Paragraph("• Scheduled Time Slots", bullet_style)],
    ]
    t = Table(data, colWidths=[350, 350])
    t.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TEXTCOLOR', (0,0), (-1,-1), colors.white),
        ('LINEBELOW', (0,0), (-1,0), 1, colors.HexColor("#D4A853")),
        ('BOTTOMPADDING', (0,0), (-1,-1), 20),
    ]))
    story.append(t)
    story.append(PageBreak())

    # Slide 5: Proven Results
    story.append(Paragraph("Proven Impact", title_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("<b>85%</b> Graduate Employment Rate", 
                           ParagraphStyle('BigStat', parent=body_style, fontSize=30, textColor=colors.HexColor("#4ADE80"))))
    story.append(Paragraph("Within 6 months of graduation", body_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>150%</b> Salary Increase", 
                           ParagraphStyle('BigStat2', parent=body_style, fontSize=30, textColor=colors.HexColor("#4ADE80"))))
    story.append(Paragraph("Compared to non-program graduates", body_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>500+</b> Students Trained Annually", 
                           ParagraphStyle('BigStat3', parent=body_style, fontSize=30, textColor=colors.HexColor("#4ADE80"))))
    story.append(PageBreak())

    # Slide 6: Investment
    story.append(Paragraph("Simple, Transparent Investment", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>$500 / month</b>", 
                           ParagraphStyle('Price', parent=title_style, fontSize=48, textColor=colors.HexColor("#D4A853"))))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Per School / Cohort", ParagraphStyle('Center', parent=body_style, alignment=TA_CENTER)))
    story.append(Spacer(1, 30))
    story.append(Paragraph("Includes everything: Internet, Staff, Tools, Curriculum.", body_style))
    story.append(Paragraph("No hidden fees. Pay-as-you-go monthly.", body_style))
    story.append(PageBreak())

    # Slide 7: Next Steps
    story.append(Spacer(1, 100))
    story.append(Paragraph("Let's Build the Future Together", title_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("Schedule your strategy consultation today.", 
                           ParagraphStyle('EndBody', parent=body_style, alignment=TA_CENTER)))
    story.append(Spacer(1, 20))
    story.append(Paragraph("www.dignidigital.com", 
                           ParagraphStyle('Link', parent=body_style, alignment=TA_CENTER, textColor=colors.HexColor("#D4A853"))))

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    print("Presentation created successfully.")

# --- Word Proposal Generation ---
def create_proposal():
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Future-Ready Graduate Program Partnership', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Prepared for: [School Name]').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('By: Digni Digital').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "Digni Digital proposes a strategic partnership with [School Name] to implement the 'Future-Ready Graduate' program. "
        "This comprehensive employability initiative is designed to equip your final-year students with high-demand digital skills, "
        "ensuring they graduate not just with a certificate, but with a career path. Our turnkey solution addresses the critical "
        "skills gap by providing expert facilitation, internet connectivity, and industry-aligned curriculum within your school premises."
    )

    # Problem Statement
    doc.add_heading('2. The Problem', level=1)
    doc.add_paragraph(
        "The traditional education model is struggling to keep pace with the rapidly evolving digital economy. "
        "Current statistics indicate a 40% unemployment rate for fresh graduates within the first six months. "
        "Employers consistently report that graduates lack practical, job-ready skills such as digital literacy, "
        "remote collaboration, and proficiency with modern AI tools."
    )

    # Proposed Solution
    doc.add_heading('3. Proposed Solution', level=1)
    doc.add_paragraph(
        "The Future-Ready Graduate Program is a full academic year curriculum integrated seamlessly into your existing schedule. "
        "We focus on three core trimesters:"
    )
    doc.add_paragraph("• Trimester 1: Digital Foundation & Web Development", style='List Bullet')
    doc.add_paragraph("• Trimester 2: Digital Marketing & Professional Branding", style='List Bullet')
    doc.add_paragraph("• Trimester 3: Job Readiness & Industry Placement", style='List Bullet')

    # Scope of Services
    doc.add_heading('4. Scope of Services', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Digni Digital Responsibilities'
    hdr_cells[1].text = 'School Responsibilities'

    row_cells = table.add_row().cells
    row_cells[0].text = "• Provide dedicated on-site facilitator\n• Install and maintain high-speed internet\n• Provide premium AI tool licenses\n• Deliver curriculum and assessments"
    row_cells[1].text = "• Provide computer lab space\n• Ensure reliable electricity supply\n• Provide student computers\n• Allocate timetable slots"

    # Investment
    doc.add_heading('5. Investment', level=1)
    doc.add_paragraph("We offer a simple, all-inclusive monthly subscription model for the partnership.")
    
    p = doc.add_paragraph()
    run = p.add_run("$500 USD / Month")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    doc.add_paragraph("This fee covers all facilitator costs, internet connectivity, software licenses, and administrative support.")

    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "By partnering with Digni Digital, [School Name] will position itself as a forward-thinking institution "
        "that prioritizes student outcomes. We look forward to transforming the future of your graduates together."
    )

    doc.save('public/docs/Proposal_FutureReadyGraduate.docx')
    print("Proposal created successfully.")

# --- Contract Generation ---
def create_contract():
    doc = Document()
    
    doc.add_heading('SERVICE LEVEL AGREEMENT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "This Service Level Agreement ('Agreement') is made and entered into on [Date], by and between:"
    )
    
    doc.add_paragraph("Digni Digital (hereinafter referred to as 'Service Provider')", style='List Bullet')
    doc.add_paragraph("AND", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[School Name] (hereinafter referred to as 'Client')", style='List Bullet')

    doc.add_heading('1. SCOPE OF WORK', level=2)
    doc.add_paragraph(
        "The Service Provider agrees to deliver the 'Future-Ready Graduate Program' as detailed in the attached Proposal. "
        "Key deliverables include: \n"
        "a) Deployment of one qualified on-site facilitator.\n"
        "b) Provision of high-speed internet connectivity for the computer lab.\n"
        "c) Delivery of the approved digital skills curriculum.\n"
        "d) Access to necessary AI tools and software."
    )

    doc.add_heading('2. TERM', level=2)
    doc.add_paragraph(
        "This Agreement shall commence on [Start Date] and continue for a period of one academic year (approx. 10 months), "
        "ending on [End Date], unless terminated earlier in accordance with this Agreement."
    )

    doc.add_heading('3. FEES AND PAYMENT', level=2)
    doc.add_paragraph(
        "The Client agrees to pay the Service Provider a monthly fee of $500 USD.\n"
        "a) Invoices will be issued on the 1st of each month.\n"
        "b) Payment is due within 7 days of invoice date.\n"
        "c) Late payments may incur a 5% penalty."
    )

    doc.add_heading('4. CLIENT OBLIGATIONS', level=2)
    doc.add_paragraph(
        "The Client agrees to:\n"
        "a) Provide a suitable computer lab environment with power.\n"
        "b) Ensure students have access to functional computers.\n"
        "c) Integrate the program into the school timetable."
    )

    doc.add_heading('5. CONFIDENTIALITY', level=2)
    doc.add_paragraph(
        "Both parties agree to keep confidential all proprietary information exchanged during the course of this Agreement."
    )

    doc.add_heading('6. TERMINATION', level=2)
    doc.add_paragraph(
        "Either party may terminate this Agreement with 30 days' written notice. In the event of termination, "
        "the Client is liable for payment of services rendered up to the effective termination date."
    )

    doc.add_paragraph("\n\n")
    
    # Signatures
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    row = table.rows[0]
    
    c1 = row.cells[0]
    c1.text = "For: Digni Digital\n\nName: ___________________\n\nTitle: ____________________\n\nDate: ____________________"
    
    c2 = row.cells[1]
    c2.text = "For: [School Name]\n\nName: ___________________\n\nTitle: ____________________\n\nDate: ____________________"

    doc.save('public/docs/Contract_FutureReadyGraduate.docx')
    print("Contract created successfully.")

if __name__ == "__main__":
    try:
        create_presentation()
        create_proposal()
        create_contract()
        print("All documents generated successfully.")
    except Exception as e:
        print(f"Error generating documents: {e}")
